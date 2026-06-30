
from __future__ import annotations
from datetime import datetime

NAVY = "1F2A44"


def _latest_status(led: dict) -> dict:
    latest: dict = {}
    for s in led.get("status_history", []):
        if s.get("workstream"):
            latest[s["workstream"]] = s
    return latest


def export_excel(led: dict, path: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    def sheet(ws, headers, rows):
        ws.append(headers)
        for c in ws[1]:
            c.font = Font(bold=True, color="FFFFFF", name="Arial")
            c.fill = PatternFill("solid", fgColor=NAVY)
            c.alignment = Alignment(vertical="center")
        for r in rows:
            ws.append(r)
        for col_i, header in enumerate(headers, 1):
            longest = max([len(str(header))] +
                          [len(str(r[col_i - 1])) for r in rows]) if rows else len(str(header))
            ws.column_dimensions[get_column_letter(col_i)].width = min(max(longest + 2, 10), 60)
        for row in ws.iter_rows(min_row=2):
            for c in row:
                c.font = Font(name="Arial")
                c.alignment = Alignment(vertical="top", wrap_text=True)
        ws.freeze_panes = "A2"
        if rows:
            ws.auto_filter.ref = ws.dimensions

    ws = wb.active
    ws.title = "Action Tracker"
    sheet(ws, ["ID", "Status", "Description", "Owner", "Workstream", "Due", "Priority", "History"],
          [[a.get("gid"), a.get("status"), a.get("description"), a.get("owner"),
            a.get("workstream"), a.get("due_date"), a.get("priority"),
            " → ".join(f'{h.get("date")}:{h.get("status")}' for h in a.get("history", []))]
           for a in led.get("tracked_actions", [])])

    sheet(wb.create_sheet("Risk Register"),
          ["ID", "Category", "Compliance", "Likelihood", "Impact", "Description", "Mitigation"],
          [[r.get("gid"), r.get("category"), "YES" if r.get("compliance_flag") else "",
            r.get("likelihood"), r.get("impact"), r.get("description"), r.get("mitigation")]
           for r in led.get("tracked_risks", [])])

    sheet(wb.create_sheet("Decision Log"),
          ["Date", "Decision", "Decided by", "Compliance", "Rationale"],
          [[d.get("meeting_date"), d.get("description"), d.get("decided_by"),
            "YES" if d.get("compliance_flag") else "", d.get("rationale")]
           for d in led.get("decisions_log", [])])

    sheet(wb.create_sheet("Status"),
          ["Workstream", "RAG", "Summary", "Blockers"],
          [[ws_name, (s.get("rag_status") or "").upper(), s.get("summary"),
            ", ".join(s.get("blockers", []))]
           for ws_name, s in _latest_status(led).items()])

    wb.save(path)
    return path


def export_word(led: dict, path: str) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    doc.add_heading(f'{led.get("project_name", "Project")} — Status Report', level=0)
    doc.add_paragraph(datetime.now().strftime("Generated %d %B %Y"))

    doc.add_heading("Delivery status by workstream", level=1)
    latest = _latest_status(led)
    if latest:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text = "Workstream", "RAG", "Summary"
        for ws_name, s in latest.items():
            c = t.add_row().cells
            c[0].text = ws_name
            c[1].text = (s.get("rag_status") or "").upper()
            c[2].text = s.get("summary") or ""
    else:
        doc.add_paragraph("No status recorded yet.")

    doc.add_heading("Open actions", level=1)
    opens = [a for a in led.get("tracked_actions", []) if a.get("status") != "done"]
    if opens:
        for a in opens:
            doc.add_paragraph(
                f'{a.get("gid")} — {a.get("description")} '
                f'(owner: {a.get("owner") or "unassigned"}, due: {a.get("due_date") or "—"}, '
                f'status: {a.get("status")})', style="List Bullet")
    else:
        doc.add_paragraph("None open.")

    doc.add_heading("Risks requiring attention", level=1)
    risks = led.get("tracked_risks", [])
    if risks:
        for r in risks:
            tag = "  [COMPLIANCE]" if r.get("compliance_flag") else ""
            doc.add_paragraph(
                f'{r.get("gid")} — {r.get("description")} '
                f'(likelihood {r.get("likelihood")}, impact {r.get("impact")}){tag}',
                style="List Bullet")
    else:
        doc.add_paragraph("None recorded.")

    doc.add_heading("Decisions", level=1)
    decisions = led.get("decisions_log", [])
    if decisions:
        for d in decisions:
            tag = "  [COMPLIANCE]" if d.get("compliance_flag") else ""
            doc.add_paragraph(
                f'({d.get("meeting_date")}) {d.get("description")} '
                f'— {d.get("decided_by") or "unattributed"}{tag}', style="List Bullet")
    else:
        doc.add_paragraph("None recorded.")

    doc.save(path)
    return path
